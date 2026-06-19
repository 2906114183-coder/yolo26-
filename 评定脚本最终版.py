import sys
import os
import gc
from pathlib import Path
import cv2
import torch
import numpy as np
from ultralytics import YOLO
from collections import Counter
import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


ROTATE_ADAPTATION_ENABLED = True  # 功能开关: 自动旋转竖向图片
EXPERT_REVIEW_ENABLED = True  # 【核心功能】功能开关: 专家知识后处理审查模块
GENERATE_WORD_REPORT_ENABLED = True  # 功能开关: 是否生成Word版报告


MIN_AREA_IN_WELD_ZONE = 0.8  # 检测框必须至少有80%的面积在焊缝区内
CONTEXT_GRADIENT_THRESHOLD = 1.8  # 上下文区域的梯度阈值，低于此值可能为焊缝外噪声
# 2. 线性缺陷智能桥接参数
LINEAR_DEFECT_BRIDGE_GAP_MAX = 80  # 线性缺陷中断的最大允许距离（像素）
LINEAR_DEFECT_ALIGNMENT_TOLERANCE = 20  # 线性缺陷对齐的容忍度（像素）
BRIDGE_PROFILE_THRESHOLD_RATIO = 0.7  # “桥梁”区域像素剖面的凹陷/凸起程度需达到缺陷本身的70%

# 修正：按训练时的类别顺序调整阈值（SL在前，Pore在后）
PER_CLASS_THRESHOLDS = {'SL': 0.32, 'Pore': 0.40, 'Crack': 0.64, 'LP': 0.40, 'LF': 0.35, 'default': 0.1}
CONTAINMENT_RECLASSIFICATION_ENABLED = True
POSITIONAL_AWARE_SUPPRESSION_ENABLED = True
MIDDLE_ZONE_RATIO = 0.6
TEST_TIME_AUGMENTATION_ENABLED = True
TTA_CONFIG = {'horizontal_flip': True, 'brightness_adjust': [0.8, 1.25], 'gamma_adjust': [0.8, 1.4]}
SHAPE_RECLASSIFICATION_ENABLED = True
CONFIDENCE_RECLASSIFICATION_ENABLED = True
PORE_LIKE_ASPECT_RATIO_MIN = 0.70
PORE_LIKE_ASPECT_RATIO_MAX = 1.40
CROSS_CLASS_SUPPRESSION_ENABLED = True
# 修正：SUPPRESSION_RULES中的顺序也对应调整（SL在前，Pore在后）
SUPPRESSION_RULES = [('SL', 'Pore', 0.45)]
INTELLIGENT_ENHANCEMENT_ENABLED = True
INTELLIGENT_ENHANCE_ALPHAS = [0.8, 1.0, 1.3, 1.6, 2.0]
ZOOM_AND_SPLIT_ENABLED = True
ZOOM_SPLIT_COUNT = 5
ZOOM_SPLIT_OVERLAP_RATIO = 0.2
POST_MERGE_ENABLED = True
MERGE_MAX_DISTANCE_X, MERGE_MAX_DISTANCE_Y = 30, 15
FINAL_NMS_IOU_THRESHOLD = 0.45
WELD_ZONE_FILTER_ENABLED = True  # 这个开关由专家模块内部的逻辑接管，但保留以备用
WELD_ZONE_HEIGHT_RATIO = 0.7

# 核心修正：CLASS_NAMES顺序改为训练时的 ["SL", "Pore", "Crack", "LP", "LF"]
CLASS_NAMES = ['SL', 'Pore', 'Crack', 'LP', 'LF']
CLASS_COLORS = [(0, 165, 255), (0, 255, 0), (0, 0, 255), (255, 0, 0), (255, 0, 255)]
INPUT_FOLDER_NAME, OUTPUT_DIR_NAME = 'D:\yolo26\输入', '输出'
BOX_THICKNESS, ARROW_THICKNESS, LABEL_PADDING = 2, 2, 10
LABEL_FONT_SCALE, LABEL_FONT_THICKNESS = 0.7, 2
LABEL_TEXT_COLOR, LABEL_BACKGROUND_COLOR = (0, 0, 0), (255, 255, 255)
VERTICAL_SHIFT_ON_OVERLAP = 30
SUMMARY_ENABLED = True



def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(os.path.dirname(__file__))
    return os.path.join(base_path, relative_path)


def expert_review_module(detections, image, grad_mag, class_names):

    print("  - Running Expert Review Module...")
    h, w = image.shape[:2]

    # --- 1. 上下文与严格ROI验证 ---
    print("    - Step 1: Context & Strict ROI Verification...")
    context_verified_dets = []
    y_start, y_end = int(h * (1 - WELD_ZONE_HEIGHT_RATIO) / 2), int(h * (1 + WELD_ZONE_HEIGHT_RATIO) / 2)
    weld_zone = (0, y_start, w, y_end)

    for det in detections:
        box = det['box']
        # 计算IoA (Intersection over Area of the box)
        inter_x1, inter_y1 = max(box[0], weld_zone[0]), max(box[1], weld_zone[1])
        inter_x2, inter_y2 = min(box[2], weld_zone[2]), min(box[3], weld_zone[3])
        inter_area = max(0, inter_x2 - inter_x1) * max(0, inter_y2 - inter_y1)
        box_area = (box[2] - box[0]) * (box[3] - box[1])
        ioa = inter_area / box_area if box_area > 0 else 0

        if ioa < MIN_AREA_IN_WELD_ZONE:
            continue

        ctx_x1, ctx_y1 = max(0, int(box[0] - (box[2] - box[0]) * 0.5)), max(0, int(box[1] - (box[3] - box[1]) * 0.5))
        ctx_x2, ctx_y2 = min(w, int(box[2] + (box[2] - box[0]) * 0.5)), min(h, int(box[3] + (box[3] - box[1]) * 0.5))
        context_patch = grad_mag[ctx_y1:ctx_y2, ctx_x1:ctx_x2]
        context_avg_grad = np.mean(context_patch) if context_patch.size > 0 else 0

        if context_avg_grad > CONTEXT_GRADIENT_THRESHOLD:
            context_verified_dets.append(det)

    # --- 2. 线性缺陷智能桥接 ---
    print("    - Step 2: Linear Defect Smart Bridging...")
    linear_defect_ids = {class_names.index(name) for name in ['Crack', 'LP', 'LF'] if name in class_names}
    linear_dets = [d for d in context_verified_dets if d['cls_id'] in linear_defect_ids]
    other_dets = [d for d in context_verified_dets if d['cls_id'] not in linear_defect_ids]

    if len(linear_dets) < 2:
        return context_verified_dets

    gray_image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY) if len(image.shape) > 2 else image

    while True:
        merged_in_pass = False
        linear_dets.sort(key=lambda d: (d['cls_id'], d['box'][0]))
        merged, used = [], set()

        for i in range(len(linear_dets)):
            if i in used: continue
            current_det = linear_dets[i].copy()

            for j in range(i + 1, len(linear_dets)):
                if j in used or current_det['cls_id'] != linear_dets[j]['cls_id']: continue
                other_det = linear_dets[j]
                c_box, o_box = current_det['box'], other_det['box']
                c_y_center, o_y_center = (c_box[1] + c_box[3]) / 2, (o_box[1] + o_box[3]) / 2
                x_gap = o_box[0] - c_box[2]

                if 0 < x_gap < LINEAR_DEFECT_BRIDGE_GAP_MAX and abs(
                        c_y_center - o_y_center) < LINEAR_DEFECT_ALIGNMENT_TOLERANCE:
                    bridge_y = int((c_y_center + o_y_center) / 2)
                    bridge_x_start, bridge_x_end = int(c_box[2]), int(o_box[0])

                    if bridge_y >= gray_image.shape[0] or bridge_x_start >= bridge_x_end:
                        continue

                    pixel_profile = gray_image[bridge_y, bridge_x_start:bridge_x_end]

                    c_patch = gray_image[int(c_box[1]):int(c_box[3]), int(c_box[0]):int(c_box[2])]
                    o_patch = gray_image[int(o_box[1]):int(o_box[3]), int(o_box[0]):int(o_box[2])]
                    defect_mean_intensity = (np.mean(c_patch) + np.mean(
                        o_patch)) / 2 if c_patch.size > 0 and o_patch.size > 0 else 0
                    bridge_mean_intensity = np.mean(pixel_profile) if pixel_profile.size > 0 else defect_mean_intensity

                    if defect_mean_intensity > 0 and bridge_mean_intensity < defect_mean_intensity / BRIDGE_PROFILE_THRESHOLD_RATIO:
                        current_det['box'] = [min(c_box[0], o_box[0]), min(c_box[1], o_box[1]), max(c_box[2], o_box[2]),
                                              max(c_box[3], o_box[3])]
                        current_det['conf'] = max(current_det['conf'], other_det['conf'])
                        used.add(j);
                        merged_in_pass = True;
                        break

            merged.append(current_det);
            used.add(i)

        linear_dets = merged
        if not merged_in_pass: break

    return other_dets + linear_dets


# --- 您的v46原有辅助函数 ---
def create_word_report(report_data, output_dir, output_filename="detection_report.docx"):
    if not DOCX_AVAILABLE:
        print("\n!!! 'python-docx' 库未找到。Word报告生成功能已禁用。 !!!");
        print("!!! 请在命令行运行 'pip install python-docx' 来启用此功能。 !!!");
        return
    doc = Document()
    title = doc.add_heading('焊缝缺陷检测综合报告', level=0);
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph();
    p.add_run(f"报告生成时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True;
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    for result in report_data:
        image_name, image_path, detections = result[
            'image_name'], output_dir / f"{Path(result['image_name']).stem}.jpg", result['detections']
        doc.add_heading(f"图片: {image_name}", level=2)
        try:
            doc.add_picture(str(image_path), width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except FileNotFoundError:
            doc.add_paragraph(f"[错误: 无法找到图片 '{image_path}']")
        if not detections:
            doc.add_paragraph("  >> 此图片中未检测到缺陷。")
        else:
            table = doc.add_table(rows=1, cols=4);
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '缺陷序号';
            hdr_cells[1].text = '性质';
            hdr_cells[2].text = '位置 (x1, y1, x2, y2)';
            hdr_cells[3].text = '置信度'
            sorted_dets = sorted(detections, key=lambda x: x['box'][0])
            for i, det in enumerate(sorted_dets):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i + 1);
                row_cells[1].text = CLASS_NAMES[det['cls_id']];
                row_cells[2].text = f"({', '.join(map(str, map(int, det['box'])))})";
                row_cells[3].text = f"{det['conf']:.2f}"
        doc.add_page_break()
    report_path = output_dir / output_filename;
    doc.save(report_path);
    print(f"Word版详细报告已生成: {report_path.resolve()}")


def de_rotate_box(box, h): x1, y1, x2, y2 = box;nx1, ny1 = y1, h - x2;nx2, ny2 = y2, h - x1;return [min(nx1, nx2),
                                                                                                    min(ny1, ny2),
                                                                                                    max(nx1, nx2),
                                                                                                    max(ny1, ny2)]


def apply_gamma_correction(img, g): inv_g = 1.0 / g;tbl = np.array(
    [((i / 255.0) ** inv_g) * 255 for i in np.arange(0, 256)]).astype("uint8");return cv2.LUT(img, tbl)


def deaugment_box(box, w, f): return [w - box[2], box[1], w - box[0], box[3]] if f else box


def is_box_inside(ib, ob): return ob[0] <= ib[0] and ob[1] <= ib[1] and ob[2] >= ib[2] and ob[3] >= ib[3]


def apply_containment_reclassification(dets, names):
    try:
        # 修正：按新的类别顺序获取SL和Pore的id（SL在前，Pore在后）
        sl_id, p_id = names.index('SL'), names.index('Pore')
    except:
        return dets
    # 修正：变量名对应调整，逻辑不变
    sl_dets, p_dets, to_rec = [(i, d) for i, d in enumerate(dets) if d['cls_id'] == sl_id], [d for d in dets if
                                                                                             d['cls_id'] == p_id], set()
    for i, sl in sl_dets:
        for p in p_dets:
            if is_box_inside(sl['box'], p['box']): to_rec.add(i);break
    for i in to_rec: dets[i]['cls_id'] = p_id
    return dets


def apply_positional_suppression(dets, w, r, names):
    try:
        lp_id, lf_id = names.index('LP'), names.index('LF')
    except:
        return dets
    to_rem, lp_dets, lf_dets = set(), [(i, d) for i, d in enumerate(dets) if d['cls_id'] == lp_id], [(i, d) for i, d in
                                                                                                     enumerate(dets) if
                                                                                                     d[
                                                                                                         'cls_id'] == lf_id]
    if not lp_dets or not lf_dets: return dets
    s, e = w * (1 - r) / 2, w * (1 + r) / 2
    for i, lp in lp_dets:
        for j, lf in lf_dets:
            if i in to_rem or j in to_rem: continue
            b1, b2 = lp['box'], lf['box']
            if max(b1[0], b2[0]) < min(b1[2], b2[2]) and max(b1[1], b2[1]) < min(b1[3], b2[3]):
                cx = (max(b1[0], b2[0]) + min(b1[2], b2[2])) / 2
                if s <= cx <= e:
                    to_rem.add(j)
                else:
                    to_rem.add(i)
    return [d for i, d in enumerate(dets) if i not in to_rem]


def apply_shape_based_reclassification(dets, names, min_ar, max_ar):
    try:
        # 修正：按新的类别顺序获取SL和Pore的id（SL在前，Pore在后）
        sl_id, p_id = names.index('SL'), names.index('Pore')
    except:
        return dets
    for d in dets:
        # 修正：判断SL类别的id
        if d['cls_id'] == sl_id:
            w, h = d['box'][2] - d['box'][0], d['box'][3] - d['box'][1]
            if h > 0 and min_ar <= (w / h) <= max_ar: d['cls_id'] = p_id
    return dets


def apply_confidence_based_reclassification(dets, names):
    try:
        # 修正：按新的类别顺序获取SL和Pore的id（SL在前，Pore在后）
        sl_id, p_id = names.index('SL'), names.index('Pore')
    except:
        return dets
    for d in dets:
        # 修正：判断SL类别的id
        if d['cls_id'] == sl_id and d['conf'] < 0.8: d['cls_id'] = p_id
    return dets


def merge_nearby_boxes(dets, max_x, max_y):
    if len(dets) < 2: return dets
    while True:
        mp = False;
        dets.sort(key=lambda d: (d['cls_id'], d['box'][0]));
        m, u = [], set()
        for i in range(len(dets)):
            if i in u: continue
            curr = dets[i].copy()
            for j in range(i + 1, len(dets)):
                if j in u or curr['cls_id'] != dets[j]['cls_id']: continue
                cb, ob = curr['box'], dets[j]['box']
                if max(0, max(cb[0], ob[0]) - min(cb[2], ob[2])) <= max_x and max(0, max(cb[1], ob[1]) - min(cb[3], ob[
                    3])) <= max_y:
                    curr['box'] = [min(cb[0], ob[0]), min(cb[1], ob[1]), max(cb[2], ob[2]), max(cb[3], ob[3])]
                    curr['conf'] = max(curr['conf'], dets[j]['conf']);
                    u.add(j);
                    mp = True
            m.append(curr);
            u.add(i)
        dets = m
        if not mp: break
    return dets


def find_optimal_enhancement(model, img, alphas, dev):
    bs, bi = -1, img
    for a in alphas:
        enh = img if a == 1.0 else cv2.convertScaleAbs(img, alpha=a, beta=10)
        res = model.predict(source=enh, conf=0.1, device=dev, verbose=False)
        s = torch.sum(res[0].boxes.conf).item() if res and res[0].boxes else 0
        if s > bs: bs, bi = s, enh
    return bi


def non_max_suppression(boxes, scores, classes, iou_t):
    if len(boxes) == 0: return []
    idxs = np.argsort(scores)[::-1];
    keep = []
    while len(idxs) > 0:
        i = idxs[0];
        keep.append(i)
        if len(idxs) == 1: break
        xx1, yy1 = np.maximum(boxes[i, 0], boxes[idxs[1:], 0]), np.maximum(boxes[i, 1], boxes[idxs[1:], 1])
        xx2, yy2 = np.minimum(boxes[i, 2], boxes[idxs[1:], 2]), np.minimum(boxes[i, 3], boxes[idxs[1:], 3])
        w, h = np.maximum(0.0, xx2 - xx1), np.maximum(0.0, yy2 - yy1);
        inter = w * h
        area_i, area_o = (boxes[i, 2] - boxes[i, 0]) * (boxes[i, 3] - boxes[i, 1]), (
                    boxes[idxs[1:], 2] - boxes[idxs[1:], 0]) * (boxes[idxs[1:], 3] - boxes[idxs[1:], 1])
        iou = inter / (area_i + area_o - inter + 1e-8)
        rem = np.where((iou <= iou_t) | (classes[i] != classes[idxs[1:]]))[0]
        idxs = idxs[rem + 1]
    return keep

if __name__ == "__main__":
    if getattr(sys, 'frozen', False):
        application_path = Path(sys.executable).parent
    else:
        application_path = Path(os.path.abspath(os.path.dirname(__file__)))
    IMAGE_SOURCE_DIR, output_dir = application_path / INPUT_FOLDER_NAME, application_path / OUTPUT_DIR_NAME
    print(f"--- NDT 预测脚本 (v46 + 专家审查模块) ---")
    device = 'cuda' if torch.cuda.is_available() else 'cpu';
    print(f"使用设备: {device}")
    try:
        model = YOLO(resource_path(os.path.join('models1', 'best.pt')));print("模型加载成功!")
    except Exception as e:
        print(f"!!! 模型加载失败: {e} !!!");input("按 Enter键 退出...");sys.exit(1)
    output_dir.mkdir(parents=True, exist_ok=True)
    image_files = sorted(
        [p for p in IMAGE_SOURCE_DIR.rglob('*') if p.suffix.lower() in ['.jpg', '.jpeg', '.png', '.bmp']])
    if not image_files: print(f"!!! 在 {IMAGE_SOURCE_DIR.resolve()} 中未找到任何图片 !!!");input(
        "按 Enter键 退出...");sys.exit(1)
    print(f"\n找到 {len(image_files)} 张图片，准备进行预测...")
    all_results_for_report, total_images_processed = [], 0
    for img_path_obj in image_files:
        print(f"\n--- [ {total_images_processed + 1}/{len(image_files)} ] 处理图片: {img_path_obj.name} ---")
        try:
            frame_orig = cv2.imread(str(img_path_obj))
            if frame_orig is None: print(f"警告: 无法读取图片 {img_path_obj.name}，跳过。");continue
            original_height, original_width = frame_orig.shape[:2]
            was_rotated = False
            if ROTATE_ADAPTATION_ENABLED and original_height > original_width * 1.1:
                print("检测到竖向图片，自动旋转90度进行处理...");
                frame_proc = cv2.rotate(frame_orig, cv2.ROTATE_90_CLOCKWISE);
                was_rotated = True
            else:
                frame_proc = frame_orig
            proc_height, proc_width = frame_proc.shape[:2]

            base_image = find_optimal_enhancement(model, frame_proc, INTELLIGENT_ENHANCE_ALPHAS,
                                                  device) if INTELLIGENT_ENHANCEMENT_ENABLED else frame_proc.copy()
            inference_tasks = [{'name': 'full_image', 'image': base_image, 'offset_xy': (0, 0)}]
            if ZOOM_AND_SPLIT_ENABLED:
                num_o = max(0, ZOOM_SPLIT_COUNT - 1);
                base_w = proc_width / (ZOOM_SPLIT_COUNT - ZOOM_SPLIT_OVERLAP_RATIO * num_o) if num_o > 0 else float(
                    proc_width)
                overlap_px = int(base_w * ZOOM_SPLIT_OVERLAP_RATIO);
                x_start, i = 0, 0
                while x_start < proc_width:
                    x_end = min(int(x_start + base_w), proc_width)
                    if x_end > x_start: inference_tasks.append(
                        {'name': f'split_{i + 1}', 'image': base_image[:, x_start:x_end],
                         'offset_xy': (x_start, 0)});i += 1
                    if x_end >= proc_width: break
                    x_start = x_end - overlap_px
            all_raw_detections = []
            for task in inference_tasks:
                task_img, offset, w = task['image'], task['offset_xy'], task['image'].shape[1];
                images_to_process = [{'img': task_img, 'flipped': False}]
                if TEST_TIME_AUGMENTATION_ENABLED:
                    if TTA_CONFIG.get('horizontal_flip'): images_to_process.append(
                        {'img': cv2.flip(task_img, 1), 'flipped': True})
                    for f in TTA_CONFIG.get('brightness_adjust', []): images_to_process.append(
                        {'img': cv2.convertScaleAbs(task_img, alpha=f, beta=0), 'flipped': False})
                    for g in TTA_CONFIG.get('gamma_adjust', []): images_to_process.append(
                        {'img': apply_gamma_correction(task_img, g), 'flipped': False})
                for item in images_to_process:
                    results = model.predict(source=item['img'], conf=0.01, device=device, verbose=False)
                    if results and results[0].boxes:
                        for i in range(len(results[0].boxes)):
                            box, conf, cls_id = results[0].boxes.xyxy[i].cpu().numpy(), results[0].boxes.conf[
                                i].cpu().item(), int(results[0].boxes.cls[i].cpu().item())
                            deaug_box = deaugment_box(box, w, item['flipped'])
                            all_raw_detections.append({'box': [deaug_box[0] + offset[0], deaug_box[1] + offset[1],
                                                               deaug_box[2] + offset[0], deaug_box[3] + offset[1]],
                                                       'conf': conf, 'cls_id': cls_id})

            detections = [d for d in all_raw_detections if
                          d['conf'] >= PER_CLASS_THRESHOLDS.get(CLASS_NAMES[d['cls_id']], 1.0)]
            if CROSS_CLASS_SUPPRESSION_ENABLED and len(detections) > 1:
                to_suppress = set();
                detections.sort(key=lambda x: x['conf'], reverse=True)
                for i in range(len(detections)):
                    if i in to_suppress: continue
                    for j in range(i + 1, len(detections)):
                        if j in to_suppress: continue
                        d1, d2 = detections[i], detections[j];
                        c1, c2 = CLASS_NAMES[d1['cls_id']], CLASS_NAMES[d2['cls_id']]
                        for r1, r2, iou_t in SUPPRESSION_RULES:
                            if {c1, c2} == {r1, r2}:
                                b1, b2 = d1['box'], d2['box'];
                                ix1, iy1, ix2, iy2 = max(b1[0], b2[0]), max(b1[1], b2[1]), min(b1[2], b2[2]), min(b1[3],
                                                                                                                  b2[3])
                                inter = max(0, ix2 - ix1) * max(0, iy2 - iy1)
                                if inter > 0:
                                    union = (b1[2] - b1[0]) * (b1[3] - b1[1]) + (b2[2] - b2[0]) * (
                                                b2[3] - b2[1]) - inter
                                    if union > 0 and inter / union > iou_t: to_suppress.add(j)
                detections = [d for i, d in enumerate(detections) if i not in to_suppress]
            if detections:
                boxes, scores, classes = np.array([d['box'] for d in detections]), np.array(
                    [d['conf'] for d in detections]), np.array([d['cls_id'] for d in detections])
                final_indices = non_max_suppression(boxes, scores, classes, FINAL_NMS_IOU_THRESHOLD)
                detections = [detections[i] for i in final_indices]
            if POSITIONAL_AWARE_SUPPRESSION_ENABLED: detections = apply_positional_suppression(detections, proc_width,
                                                                                               MIDDLE_ZONE_RATIO,
                                                                                               CLASS_NAMES)
            if CONTAINMENT_RECLASSIFICATION_ENABLED: detections = apply_containment_reclassification(detections,
                                                                                                     CLASS_NAMES)
            if SHAPE_RECLASSIFICATION_ENABLED: detections = apply_shape_based_reclassification(detections, CLASS_NAMES,
                                                                                               PORE_LIKE_ASPECT_RATIO_MIN,
                                                                                               PORE_LIKE_ASPECT_RATIO_MAX)
            if CONFIDENCE_RECLASSIFICATION_ENABLED: detections = apply_confidence_based_reclassification(detections,
                                                                                                         CLASS_NAMES)
            detections_after_v46 = merge_nearby_boxes(detections, MERGE_MAX_DISTANCE_X,
                                                      MERGE_MAX_DISTANCE_Y) if POST_MERGE_ENABLED else detections

            if EXPERT_REVIEW_ENABLED:
                gray_proc = cv2.cvtColor(frame_proc, cv2.COLOR_BGR2GRAY) if len(frame_proc.shape) > 2 else frame_proc
                grad_x = cv2.Sobel(gray_proc, cv2.CV_64F, 1, 0, ksize=3);
                grad_y = cv2.Sobel(gray_proc, cv2.CV_64F, 0, 1, ksize=3)
                grad_mag = cv2.magnitude(grad_x, grad_y)
                detections_after_review = expert_review_module(detections_after_v46, frame_proc, grad_mag, CLASS_NAMES)
            else:
                detections_after_review = detections_after_v46

            final_results = []
            if was_rotated:
                for det in detections_after_review:
                    det_copy = det.copy();
                    det_copy['box'] = de_rotate_box(det['box'], original_height);
                    final_results.append(det_copy)
            else:
                final_results = detections_after_review
            all_results_for_report.append({'image_name': img_path_obj.name, 'detections': final_results})

            vis_frame = frame_orig.copy();
            final_results.sort(key=lambda d: d['box'][0]);
            label_positions_y = {}
            for det in final_results:
                x1, y1, x2, y2 = map(int, det['box']);
                name, color = CLASS_NAMES[det['cls_id']], CLASS_COLORS[det['cls_id'] % len(CLASS_COLORS)]
                cv2.rectangle(vis_frame, (x1, y1), (x2, y2), color, BOX_THICKNESS)
                label_text = f"{name} {det['conf']:.2f}";
                (tw, th), _ = cv2.getTextSize(label_text, cv2.FONT_HERSHEY_SIMPLEX, LABEL_FONT_SCALE,
                                              LABEL_FONT_THICKNESS)
                bg_w, bg_h = tw + 2 * LABEL_PADDING, th + 2 * LABEL_PADDING;
                lx, ly = x1 + (x2 - x1 - bg_w) // 2, y1 - 20 - bg_h
                if lx < 0: lx = 5
                if lx + bg_w > original_width: lx = original_width - bg_w - 5
                final_y, collided, attempts = ly, True, 0
                while collided and attempts < 20:
                    collided = False
                    for px, (py, pw, ph) in label_positions_y.items():
                        if not (lx + bg_w < px or lx > px + pw) and not (
                                final_y + bg_h < py or final_y > py + ph): collided = True;final_y -= VERTICAL_SHIFT_ON_OVERLAP;break
                    attempts += 1
                if final_y < 0: final_y = y2 + 20
                ly = final_y;
                label_positions_y[lx] = (ly, bg_w, bg_h)
                cv2.rectangle(vis_frame, (lx, ly), (lx + bg_w, ly + bg_h), LABEL_BACKGROUND_COLOR, -1);
                cv2.rectangle(vis_frame, (lx, ly), (lx + bg_w, ly + bg_h), color, 2)
                cv2.putText(vis_frame, label_text, (lx + LABEL_PADDING, ly + th + LABEL_PADDING),
                            cv2.FONT_HERSHEY_SIMPLEX, LABEL_FONT_SCALE, LABEL_TEXT_COLOR, LABEL_FONT_THICKNESS,
                            cv2.LINE_AA)
                cv2.arrowedLine(vis_frame, (lx + bg_w // 2, ly + bg_h if ly < y1 else ly),
                                (int((x1 + x2) / 2), y1 if ly < y1 else y2), color, ARROW_THICKNESS, tipLength=0.3)
            if False:
                summary_y = original_height - 120;
                cv2.putText(vis_frame, "Summary:", (20, summary_y), cv2.FONT_HERSHEY_TRIPLEX, 1.0, (220, 220, 220), 2,
                            cv2.LINE_AA)
                cx, cy = 20, summary_y + int(cv2.getTextSize("A", cv2.FONT_HERSHEY_SIMPLEX, 0.8, 2)[0][1]) + 20
                counts = Counter(CLASS_NAMES[det['cls_id']] for det in final_results)
                if not counts:
                    cv2.putText(vis_frame, "No defects detected.", (cx, cy), cv2.FONT_HERSHEY_SIMPLEX, 0.8,
                                (200, 200, 200), 2, cv2.LINE_AA)
                else:
                    for class_idx, class_name in enumerate(CLASS_NAMES):
                        count = counts.get(class_name, 0)
                        if count > 0:
                            text = f"{class_name}: {count}";
                            color = CLASS_COLORS[class_idx % len(CLASS_COLORS)];
                            (tw, th), _ = cv2.getTextSize(text, cv2.FONT_HERSHEY_SIMPLEX, 0.8, 2)
                            if cx + tw > original_width - 20: cx, cy = 20, cy + th + 20
                            cv2.putText(vis_frame, text, (cx, cy), cv2.FONT_HERSHEY_SIMPLEX, 0.8, color, 2,
                                        cv2.LINE_AA);
                            cx += tw + 40
            output_filename = output_dir / f"{img_path_obj.stem}.jpg";
            cv2.imwrite(str(output_filename), vis_frame);
            total_images_processed += 1;
            gc.collect()
        except Exception as e:
            print(f"!!! 处理图片 {img_path_obj.name} 时发生严重错误: {e} !!!");import \
                traceback;traceback.print_exc();continue

    txt_report_path = output_dir / f"detection_report_ultimate.txt"
    with open(txt_report_path, 'w', encoding='utf-8') as f:
        f.write("==================================================\n");
        f.write("     焊缝缺陷检测综合报告 (v46+专家审查模块)\n");
        f.write("==================================================\n\n")
        f.write(f"报告生成时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n");
        f.write(f"共处理图片: {total_images_processed} 张\n\n")
        if not all_results_for_report:
            f.write("未在任何图片中检测到缺陷。\n")
        else:
            for result in all_results_for_report:
                f.write(f"--------------------------------------------------\n");
                f.write(f"图片: {result['image_name']}\n");
                f.write(f"--------------------------------------------------\n")
                if not result['detections']: f.write("  >> 此图片中未检测到缺陷。\n\n");continue
                sorted_dets = sorted(result['detections'], key=lambda x: x['box'][0])
                for i, det in enumerate(sorted_dets):
                    f.write(f"  [缺陷 {i + 1}]\n")
                    x1, y1, x2, y2 = map(int, det['box']);
                    name, conf = CLASS_NAMES[det['cls_id']], det['conf']
                    f.write(f"    - 性质: {name}\n");
                    f.write(f"    - 位置 (x1, y1, x2, y2): ({x1}, {y1}, {x2}, {y2})\n");
                    f.write(f"    - 置信度: {conf:.3f}\n")
                f.write("\n")
    print(f"\n--- 所有 {total_images_processed} 张图片处理完毕 ---")
    print(f"TXT版详细报告已生成: {txt_report_path.resolve()}")

    if GENERATE_WORD_REPORT_ENABLED:
        create_word_report(all_results_for_report, output_dir)

    input("按 Enter键 退出...")